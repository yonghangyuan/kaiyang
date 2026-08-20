"""Markdown → DOCX 转换器（python-docx）。

针对《UAP 披露线 1947→2026：媒体脉络分析报告》的 Markdown 结构：
标题 / 表格 / 列表 / 引用 / 代码块 / 分割线 / 行内加粗与行内代码。
中文字体：正文 宋体，标题 微软雅黑，代码 Consolas。

用法:
    python md_to_docx.py uap-disclosure-1947-2026.md
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

EAST_BODY = "宋体"
EAST_HEAD = "微软雅黑"
LATIN_BODY = "Calibri"
LATIN_HEAD = "微软雅黑"
MONO = "Consolas"


def _set_font(style_or_run, latin: str, east: str, size: float | None = None, bold: bool | None = None, color=None):
    """设置拉丁与东亚字体。"""
    f = style_or_run.font
    f.name = latin
    rpr = f.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if color is not None:
        f.color.rgb = RGBColor(*color)


def _shade(paragraph, fill: str):
    """段落底色（代码块用）。"""
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_ITALIC = re.compile(r"(\*[^*]+\*)")


def add_rich_text(paragraph, text: str, size: float, bold_all: bool = False, color=None):
    """写入带 **加粗** / `行内代码` 的文本。"""
    for seg in _INLINE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            run = paragraph.add_run(seg[2:-2])
            _set_font(run, LATIN_BODY, EAST_BODY, size, bold=True, color=color)
        elif seg.startswith("`") and seg.endswith("`"):
            run = paragraph.add_run(seg[1:-1])
            _set_font(run, MONO, EAST_BODY, size - 0.5, bold=bold_all, color=color)
        else:
            run = paragraph.add_run(seg)
            _set_font(run, LATIN_BODY, EAST_BODY, size, bold=bold_all, color=color)


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_separator_row(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, out_path: Path | None = None) -> Path:
    out_path = out_path or md_path.with_suffix(".docx")
    doc = Document()

    # 页面与默认样式
    normal = doc.styles["Normal"]
    _set_font(normal, LATIN_BODY, EAST_BODY, 10.5)
    for level, sz, col in [(1, 18, (0x1F, 0x38, 0x64)), (2, 14, (0x1F, 0x38, 0x64)), (3, 12, (0x2E, 0x55, 0x8A))]:
        st = doc.styles[f"Heading {level}"]
        _set_font(st, LATIN_HEAD, EAST_HEAD, sz, bold=True, color=col)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(lines[i])
                _set_font(run, MONO, EAST_BODY, 8.5)
                _shade(p, "F2F2F2")
                i += 1
            i += 1
            continue

        # 表格
        if _is_table_row(line):
            rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                if not _is_separator_row(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c in range(ncols):
                        cell = table.cell(r, c)
                        text = row[c] if c < len(row) else ""
                        cell.paragraphs[0].text = ""
                        add_rich_text(cell.paragraphs[0], text, 9, bold_all=(r == 0))
                doc.add_paragraph()
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading("", level=min(level, 3))
            add_rich_text(h, m.group(2).strip(), {1: 18, 2: 14, 3: 12, 4: 11}[level], bold_all=True)
            i += 1
            continue

        # 分割线
        if re.fullmatch(r"\s*-{3,}\s*", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "BBBBBB")
            pbdr.append(bottom)
            ppr.append(pbdr)
            i += 1
            continue

        # 引用
        if line.strip().startswith(">"):
            text = line.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_rich_text(p, text, 9.5, color=(0x66, 0x66, 0x66))
            for run in p.runs:
                run.font.italic = True
            i += 1
            continue

        # 列表
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, text, 10.5)
            i += 1
            continue
        if re.match(r"^\s*\d+[.)]\s+", line):
            text = re.sub(r"^\s*\d+[.)]\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, text, 10.5)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_rich_text(p, line.strip(), 10.5)
        i += 1

    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    src = Path(sys.argv[1] if len(sys.argv) > 1 else __file__).resolve().parent / "uap-disclosure-1947-2026.md"
    out = convert(src)
    print(f"OK -> {out}")
